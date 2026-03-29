import os
import httpx
import logging
import base64
import io
import asyncio
import uuid
import zipfile
from typing import Optional

log = logging.getLogger("edict.multimodal_parser")

SILICONFLOW_API_KEY = os.getenv("SILICONFLOW_API_KEY")
SILICONFLOW_API_URL = os.getenv("SILICONFLOW_API_URL", "https://api.siliconflow.cn/v1")
MINERU_API_TOKEN = os.getenv("MINERU_API_TOKEN")
MINERU_API_BASE = "https://mineru.net"
# APP_BASE_URL: 本服务的公网访问地址，用于让 MinerU 回调下载临时文件
# 例如：http://your-server-ip:8000  或  https://your-domain.com
APP_BASE_URL = os.getenv("APP_BASE_URL", "")

# 临时文件存储路径 (会在 main.py 中被挂载为静态文件路由)
MINERU_TEMP_DIR = "/tmp/mineru_uploads"
os.makedirs(MINERU_TEMP_DIR, exist_ok=True)


class MultiModalParser:
    """多模态解析服务：负责将文档(PDF/PPT/Excel)和图片转换为 Markdown 文本。
    
    解析优先级：
    - PDF/PPT: MinerU 精准 API (VLM) → pymupdf4llm 降级
    - Excel/CSV: pandas
    - DOCX: python-docx
    - 图片: GLM VLM
    """

    def __init__(self):
        self.glm_model = "THUDM/GLM-4.1V-9B-Thinking"

    async def parse(self, file_bytes: bytes, filename: str, file_path: Optional[str] = None) -> str:
        """解析文件内容为 Markdown 字符串。
        
        Args:
            file_bytes: 原始字节流
            filename: 原文件名
            file_path: 如果文件已被 RAGService 持久化，传入物理路径以避免重复写磁盘
        """
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        
        if ext == "pdf":
            return await self._parse_with_mineru(file_bytes, filename, ext, file_path)
            
        elif ext in ("pptx", "ppt"):
            return await self._parse_with_mineru(file_bytes, filename, ext, file_path)
            
        elif ext in ("xlsx", "xls", "csv"):
            return await self._parse_with_pandas(file_bytes, ext)
            
        elif ext == "docx":
            return await self._parse_with_docx(file_bytes)
            
        elif ext in ("png", "jpg", "jpeg", "bmp", "tiff"):
            mime_type = f"image/{ext}" if ext != "jpg" else "image/jpeg"
            prompt = """
                # Role: 多模态数据提取专家
                # Task: 分析并提取上传图片的所有关键信息。
                # Constraints:
                1. 识别与分类：首先判断图片属于（纯文字文档、数据图表、自然实景、还是技术架构图）。
                2. 视觉描述：简述图片主体内容、场景。
                3. 文字识别：若有文字，必须按原排版输出 Markdown 格式，严禁漏字、错字。
                4. 数据提取：若含图表/表格，请将其转化为 Markdown 表格，并提取核心趋势或异常数值。
                5. 逻辑结构：使用清晰的分级标题组织输出。
            """
            return await self._parse_with_glm(file_bytes, prompt, mime_type)
        
        # 兜底：纯文本尝试解码
        try:
            return file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            return file_bytes.decode("gbk", errors="ignore")

    async def _parse_with_mineru(self, file_bytes: bytes, filename: str, ext: str, file_path: Optional[str] = None) -> str:
        """使用 MinerU 精准解析 API 解析 PDF/PPT。
        
        流程：
        1. 确保文件在静态资源目录下可访问（如果是已持久化的，创建软链接；否则保存临时文件）
        2. 构造公网 URL 提交 MinerU
        3. 异步轮询等待解析完成
        4. 下载结果 ZIP 中的 Markdown
        5. 清理静态资源目录下的临时项
        
        降级策略：MinerU 不可用时，PDF 降级到 pymupdf4llm。
        """
        if not MINERU_API_TOKEN:
            log.warning("MINERU_API_TOKEN not set.")
            if ext == "pdf":
                log.info("Falling back to pymupdf4llm for PDF.")
                return await self._parse_with_pymupdf4llm(file_bytes)
            return f"[解析失败: MinerU API Token 未配置，无法解析 .{ext} 文件]"
        
        if not APP_BASE_URL:
            log.warning("APP_BASE_URL not set, cannot construct public file URL for MinerU.")
            if ext == "pdf":
                return await self._parse_with_pymupdf4llm(file_bytes)
            return f"[解析失败: APP_BASE_URL 未配置]"

        # 1. 准备公开可访问的路径
        # 即使文件已经持久化在 data/uploads，也要链接到 /tmp/mineru_uploads 以便 FastAPI 静态路由访问
        temp_filename = f"{uuid.uuid4().hex}_{filename}"
        static_temp_path = os.path.join(MINERU_TEMP_DIR, temp_filename)
        
        try:
            if file_path and os.path.exists(file_path):
                # 优先尝试软连接以节省空间和 IO
                try:
                    os.symlink(os.path.abspath(file_path), static_temp_path)
                    log.info(f"MinerU: Created symlink from {file_path} to {static_temp_path}")
                except (OSError, AttributeError):
                    # 软连失败（Windows 权限等）则回写
                    with open(static_temp_path, "wb") as f:
                        f.write(file_bytes)
            else:
                with open(static_temp_path, "wb") as f:
                    f.write(file_bytes)
        except Exception as e:
            log.error(f"Failed to prepare static file for MinerU: {e}")
            return f"[预处理失败: {e}]"
        
        file_url = f"{APP_BASE_URL.rstrip('/')}/static/mineru_tmp/{temp_filename}"
        log.info(f"MinerU: Public URL for MinerU: {file_url}")
        
        try:
            result_md = await self._mineru_submit_and_poll(file_url)
            return result_md
        except Exception as e:
            log.error(f"MinerU parsing failed: {e}")
            if ext == "pdf":
                log.info("Falling back to pymupdf4llm.")
                return await self._parse_with_pymupdf4llm(file_bytes)
            return f"[MinerU 解析失败: {e}]"
        finally:
            # 清理静态映射，不影响 data/uploads 中的原件
            try:
                if os.path.lexists(static_temp_path):
                    os.remove(static_temp_path)
                    log.debug(f"MinerU static temp cleanup: {static_temp_path}")
            except OSError:
                pass

    async def _mineru_submit_and_poll(self, file_url: str, timeout: int = 300, interval: int = 5) -> str:
        """提交解析任务并轮询结果。"""
        async with httpx.AsyncClient(timeout=60.0) as client:
            # Step 1: 提交解析任务
            submit_resp = await client.post(
                f"{MINERU_API_BASE}/api/v4/extract/task",
                headers={
                    "Content-Type": "application/json",
                    "Authorization": f"Bearer {MINERU_API_TOKEN}"
                },
                json={
                    "url": file_url,
                    "model_version": "vlm"
                }
            )
            submit_resp.raise_for_status()
            submit_data = submit_resp.json()
            
            if submit_data.get("code") != 0:
                raise RuntimeError(f"MinerU submit error: {submit_data.get('msg')}")
            
            task_id = submit_data["data"]["task_id"]
            log.info(f"MinerU task submitted: {task_id}")
            
            # Step 2: 异步轮询
            elapsed = 0
            state_labels = {
                "uploading": "文件下载中",
                "pending": "排队中",
                "running": "解析中",
            }
            
            while elapsed < timeout:
                await asyncio.sleep(interval)
                elapsed += interval
                
                poll_resp = await client.get(
                    f"{MINERU_API_BASE}/api/v4/extract/task/{task_id}",
                    headers={"Authorization": f"Bearer {MINERU_API_TOKEN}"}
                )
                poll_resp.raise_for_status()
                poll_data = poll_resp.json()
                
                state = poll_data["data"].get("state", "unknown")
                
                if state == "done":
                    full_zip_url = poll_data["data"].get("full_zip_url")
                    log.info(f"MinerU task {task_id} completed in {elapsed}s. Downloading...")
                    
                    if full_zip_url:
                        return await self._download_mineru_zip(client, full_zip_url)
                    else:
                        raise RuntimeError("MinerU returned done but no full_zip_url")
                
                elif state == "failed":
                    err_msg = poll_data["data"].get("err_msg", "unknown error")
                    raise RuntimeError(f"MinerU task failed: {err_msg}")
                
                else:
                    progress = poll_data["data"].get("extract_progress", {})
                    extracted = progress.get("extracted_pages", "?")
                    total = progress.get("total_pages", "?")
                    label = state_labels.get(state, state)
                    log.info(f"MinerU [{elapsed}s] {label}... ({extracted}/{total} pages)")
            
            raise TimeoutError(f"MinerU task {task_id} timed out after {timeout}s")

    async def _download_mineru_zip(self, client: httpx.AsyncClient, zip_url: str) -> str:
        """下载 MinerU 返回的 ZIP 包并提取 Markdown 内容。"""
        resp = await client.get(zip_url, timeout=120.0)
        resp.raise_for_status()
        
        zip_buffer = io.BytesIO(resp.content)
        markdown_content = ""
        
        with zipfile.ZipFile(zip_buffer, "r") as zf:
            # 查找 .md 文件（通常名为 full.md 或与原文件同名的 .md）
            md_files = [f for f in zf.namelist() if f.endswith(".md")]
            if md_files:
                # 优先选 full.md，否则选最大的 .md 文件
                target = next((f for f in md_files if "full" in f.lower()), None)
                if not target:
                    target = max(md_files, key=lambda f: zf.getinfo(f).file_size)
                
                markdown_content = zf.read(target).decode("utf-8")
                log.info(f"MinerU: Extracted {target} ({len(markdown_content)} chars)")
            else:
                # 没有 .md 文件，尝试找 .txt
                txt_files = [f for f in zf.namelist() if f.endswith(".txt")]
                if txt_files:
                    markdown_content = zf.read(txt_files[0]).decode("utf-8")
                else:
                    raise RuntimeError(f"MinerU ZIP contains no .md or .txt files: {zf.namelist()}")
        
        return markdown_content

    # pymupdf4llm 降级方案
    async def _parse_with_pymupdf4llm(self, file_bytes: bytes) -> str:
        """使用 PyMuPDF4LLM 提取 PDF，并自动识别/处理扫描件。"""
        try:
            import fitz
            import pymupdf4llm
            doc = fitz.Document(stream=file_bytes, filetype="pdf")
            
            md_text = pymupdf4llm.to_markdown(doc)
            
            # 扫描件检测与 VLM 补偿
            if len(md_text.strip()) < 100 * doc.page_count and doc.page_count > 0:
                log.info(f"Detecting potential scanned PDF ({doc.page_count} pages). Falling back to VLM.")
                vlm_results = []
                max_pages = min(doc.page_count, 10)
                
                for i in range(max_pages):
                    page = doc.load_page(i)
                    # 将页面渲染为高清图片 (Matrix(2,2) 表示 144 DPI)
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                    img_bytes = pix.tobytes("jpg")
                    
                    prompt = f"这是文档的第 {i+1} 页。请准确提取所有文字、表格和图表内容，并以 Markdown 格式输出。如果是表格，请保持排版。"
                    page_text = await self._parse_with_glm(img_bytes, prompt, "image/jpeg")
                    vlm_results.append(f"## Page {i+1}\n\n{page_text}")
                
                return "\n\n".join(vlm_results)
                
            return md_text
        except ImportError:
            log.error("pymupdf4llm or fitz is not installed.")
            return "[解析失败: pymupdf4llm/fitz 依赖缺失]"
        except Exception as e:
            log.error(f"PyMuPDF4LLM fallback parsing error: {e}")
            return f"[PDF解析失败: {e}]"

    # 其他解析器 (pandas, docx, VLM)
    async def _parse_with_pandas(self, file_bytes: bytes, ext: str) -> str:
        """使用 pandas 提取表格为 Markdown。"""
        try:
            import pandas as pd
            if ext == "csv":
                df = pd.read_csv(io.BytesIO(file_bytes))
            else:
                df = pd.read_excel(io.BytesIO(file_bytes))
            return df.to_markdown(index=False)
        except ImportError:
             log.error("pandas or tabulate is not installed.")
             return "[解析失败: pandas 依赖缺失]"
        except Exception as e:
            log.error(f"Pandas parsing error: {e}")
            return f"[表格解析失败: {e}]"

    async def _parse_with_docx(self, file_bytes: bytes) -> str:
        """使用 python-docx 提取 Word 文档内容，按顺序保留段落和表格。"""
        try:
            from docx import Document
            from docx.table import Table
            from docx.text.paragraph import Paragraph
            
            doc = Document(io.BytesIO(file_bytes))
            md_sections = []
            
            from docx.oxml.table import CT_Tbl
            from docx.oxml.text.paragraph import CT_P

            for child in doc.element.body:
                if isinstance(child, CT_P):
                    para = Paragraph(child, doc)
                    text = para.text.strip()
                    if text:
                        style = para.style.name if para.style else ""
                        if "Heading" in style:
                            level_match = [s for s in style if s.isdigit()]
                            level = int(level_match[0]) if level_match else 1
                            md_sections.append(f"{'#' * level} {text}")
                        else:
                            md_sections.append(text)
                elif isinstance(child, CT_Tbl):
                    table = Table(child, doc)
                    table_md = []
                    for i, row in enumerate(table.rows):
                        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                        table_md.append("| " + " | ".join(cells) + " |")
                        if i == 0:
                            table_md.append("| " + " | ".join(["---"] * len(cells)) + " |")
                    md_sections.append("\n".join(table_md))
            
            if not md_sections:
                for para in doc.paragraphs:
                    if para.text.strip():
                        md_sections.append(para.text.strip())
            
            return "\n\n".join(md_sections)
            
        except ImportError:
            log.error("python-docx is not installed.")
            return "[解析失败: python-docx 依赖缺失]"
        except Exception as e:
            log.error(f"Docx parsing error: {e}")
            return f"[Word解析失败: {e}]"

    async def _call_vlm(self, model: str, prompt: str, image_b64: str, mime_type: str = "image/jpeg") -> str:
        """通用的 VLM 调用逻辑。"""
        if not SILICONFLOW_API_KEY:
            return "错误：未发现 SILICONFLOW_API_KEY，无法解析多模态文件。"

        payload = {
            "model": model,
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {
                            "type": "image_url",
                            "image_url": {"url": f"data:{mime_type};base64,{image_b64}"}
                        }
                    ]
                }
            ],
            "temperature": 0.2
        }

        try:
            async with httpx.AsyncClient() as client:
                response = await client.post(
                    f"{SILICONFLOW_API_URL}/chat/completions",
                    headers={"Authorization": f"Bearer {SILICONFLOW_API_KEY}"},
                    json=payload,
                    timeout=60.0
                )
                response.raise_for_status()
                data = response.json()
                return data["choices"][0]["message"]["content"]
        except Exception as e:
            log.error(f"VLM ({model}) call failed: {e}")
            return f"[视觉提取失败: {e}]"

    async def _parse_with_glm(self, file_bytes: bytes, prompt: str, mime_type: str = "image/jpeg") -> str:
        """使用 GLM 进行视觉深度解析。"""
        b64 = base64.b64encode(file_bytes).decode("utf-8")
        return await self._call_vlm(self.glm_model, prompt, b64, mime_type)

